"""Personal-document ingestion: PDF, DOCX, MD, TXT, TEX.

Accepts either a single file path or a folder path (recursively walked).
Returns normalized `RawDocument` list. Pure text extraction — no parsing of resume
*structure* happens here; that's the synthesizer's job.
"""

from __future__ import annotations

import logging
from pathlib import Path

from ..core.models import DocumentType, RawDocument
from .base import SourceCollector

log = logging.getLogger(__name__)

_SUPPORTED_EXT = {
    ".pdf": DocumentType.PDF,
    ".docx": DocumentType.DOCX,
    ".md": DocumentType.MD,
    ".markdown": DocumentType.MD,
    ".txt": DocumentType.TXT,
    ".tex": DocumentType.TEX,
}


class DocumentSource(SourceCollector):
    name = "document"

    def collect(self, path: str | Path) -> list[RawDocument]:
        target = Path(path).expanduser().resolve()
        if not target.exists():
            raise FileNotFoundError(f"No such path: {target}")
        files = self._gather_files(target)
        documents: list[RawDocument] = []
        for f in files:
            doc = self._read_file(f)
            if doc is not None:
                documents.append(doc)
        return documents

    @staticmethod
    def _gather_files(target: Path) -> list[Path]:
        if target.is_file():
            return [target]
        return [
            p for p in target.rglob("*")
            if p.is_file() and p.suffix.lower() in _SUPPORTED_EXT
        ]

    @classmethod
    def _read_file(cls, path: Path) -> RawDocument | None:
        suffix = path.suffix.lower()
        doc_type = _SUPPORTED_EXT.get(suffix, DocumentType.OTHER)
        try:
            if doc_type == DocumentType.PDF:
                text = cls._read_pdf(path)
            elif doc_type == DocumentType.DOCX:
                text = cls._read_docx(path)
            else:
                text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:
            log.warning("Failed to read %s: %s", path, exc)
            return None
        return RawDocument(
            path=str(path),
            filename=path.name,
            doc_type=doc_type,
            text=text,
        )

    @staticmethod
    def _read_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        from docx import Document  # python-docx

        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
